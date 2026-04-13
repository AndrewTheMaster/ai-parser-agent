"""Load and normalize text from DOCX and PDF inputs."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from app.nlp.normalize import clean_text


def _chunk_text(text: str, max_chars: int = 3500, overlap: int = 400) -> list[str]:
    text = clean_text(text)
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]
    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + max_chars, n)
        chunk = text[start:end]
        chunks.append(chunk)
        if end >= n:
            break
        start = max(0, end - overlap)
    return chunks


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n\n".join(parts))


def read_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    return clean_text("\n\n".join(parts))


def load_document(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".docx":
        return read_docx(path)
    if suf == ".pdf":
        return read_pdf_text(path)
    raise ValueError(f"Unsupported input type: {path}")


def document_to_chunks(
    path: Path,
    *,
    max_chars: int = 3500,
    overlap: int = 400,
) -> tuple[str, list[dict[str, Any]]]:
    """Return full cleaned text and overlapping chunks with stable ids."""
    full = load_document(path)
    raw_chunks = _chunk_text(full, max_chars=max_chars, overlap=overlap)
    chunks: list[dict[str, Any]] = []
    stem = path.name
    for i, c in enumerate(raw_chunks):
        cid = f"{stem}::chunk::{i}"
        chunks.append({"id": cid, "path": str(path.resolve()), "text": c, "index": i})
    return full, chunks


def merge_chunk_lists(chunk_lists: list[list[dict[str, Any]]]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for lst in chunk_lists:
        for ch in lst:
            cid = ch["id"]
            if cid not in seen:
                seen.add(cid)
                out.append(ch)
    return out


def strip_control_chars(s: str) -> str:
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s)
